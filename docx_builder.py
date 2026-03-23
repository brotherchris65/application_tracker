"""
docx_builder.py — Builds a formatted .docx from tailored resume data.
Returns bytes ready for st.download_button.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── Colour constants ──────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x38, 0x64)   # name header
BLUE_MID   = RGBColor(0x2E, 0x74, 0xB5)   # section headings & job titles
GREY       = RGBColor(0x66, 0x66, 0x66)   # contact line & dates
BLACK      = RGBColor(0x00, 0x00, 0x00)


def _set_color(run, color: RGBColor):
    run.font.color.rgb = color


def _hr(doc, color_hex="2E74B5"):
    """Add a horizontal rule via paragraph bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    _set_color(run, BLUE_MID)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    _set_color(run, BLACK)
    return p


def build(resume_data: dict, job_title: str, company: str) -> bytes:
    """Build and return .docx bytes from resume_data dict."""
    doc = Document()

    # ── Page margins (0.6 inch all round) ────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.65)
        section.right_margin  = Inches(0.65)

    # ── Default paragraph spacing ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Header: Name ──────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run("Christopher Hill")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Arial"
    _set_color(name_run, BLUE_DARK)

    # Title line
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    t_run = title_para.add_run("Data Engineering Manager  ·  BI Lead")
    t_run.font.size = Pt(11)
    t_run.font.name = "Arial"
    _set_color(t_run, BLUE_MID)

    # Contact line
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(4)
    c_run = contact_para.add_run(
        "Burleson, TX  ·  817-805-0832  ·  brotherchris65@gmail.com  ·  linkedin.com/in/chris-hill02020"
    )
    c_run.font.size = Pt(9)
    c_run.font.name = "Arial"
    _set_color(c_run, GREY)

    _hr(doc)

    # ── Summary ───────────────────────────────────────────────────────────────
    _section_heading(doc, "Professional Summary")
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(4)
    s_run = summary_para.add_run(resume_data.get("summary", ""))
    s_run.font.size = Pt(10)
    s_run.font.name = "Arial"
    _set_color(s_run, BLACK)

    _hr(doc)

    # ── Skills ────────────────────────────────────────────────────────────────
    _section_heading(doc, "Technical Skills & Certifications")
    skills = resume_data.get("skills", {})
    skill_map = [
        ("Platforms",         skills.get("platforms")),
        ("Languages",         skills.get("languages")),
        ("Data Integration",  skills.get("data_integration")),
        ("AI Development",    skills.get("ai_development")),
        ("Frameworks",        skills.get("frameworks")),
        ("Data Engineering",  skills.get("data_engineering")),
        ("Analysis",          skills.get("analysis")),
    ]
    for label, value in skill_map:
        if not value:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        bold_run = p.add_run(f"{label}: ")
        bold_run.bold = True
        bold_run.font.size = Pt(10)
        bold_run.font.name = "Arial"
        _set_color(bold_run, BLACK)
        val_run = p.add_run(value)
        val_run.font.size = Pt(10)
        val_run.font.name = "Arial"
        _set_color(val_run, BLACK)

    # Certifications
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    cert_bold = p.add_run("Certifications: ")
    cert_bold.bold = True
    cert_bold.font.size = Pt(10)
    cert_bold.font.name = "Arial"
    _set_color(cert_bold, BLACK)
    cert_val = p.add_run("Google Data Analytics Certification (2021)  ·  Google Certified Educator Level 2 (2020)")
    cert_val.font.size = Pt(10)
    cert_val.font.name = "Arial"
    _set_color(cert_val, BLACK)

    _hr(doc)

    # ── Experience ────────────────────────────────────────────────────────────
    _section_heading(doc, "Professional Experience")

    for exp in resume_data.get("experience", []):
        # Company + dates on same line (tab-separated via two runs)
        exp_para = doc.add_paragraph()
        exp_para.paragraph_format.space_before = Pt(8)
        exp_para.paragraph_format.space_after  = Pt(1)

        comp_run = exp_para.add_run(exp.get("company", ""))
        comp_run.bold = True
        comp_run.font.size = Pt(11)
        comp_run.font.name = "Arial"
        _set_color(comp_run, BLACK)

        # Spacer then dates right-aligned via tab stop
        tab_run = exp_para.add_run("  —  ")
        tab_run.font.size = Pt(10)
        tab_run.font.name = "Arial"
        _set_color(tab_run, GREY)

        date_run = exp_para.add_run(exp.get("dates", ""))
        date_run.font.size = Pt(10)
        date_run.font.name = "Arial"
        _set_color(date_run, GREY)

        # Job title
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after  = Pt(3)
        t_run = title_para.add_run(exp.get("title", ""))
        t_run.italic = True
        t_run.font.size = Pt(10)
        t_run.font.name = "Arial"
        _set_color(t_run, BLUE_MID)

        for bullet_text in exp.get("bullets", []):
            _add_bullet(doc, bullet_text)

    _hr(doc)

    # ── Education ─────────────────────────────────────────────────────────────
    _section_heading(doc, "Education")
    edu_items = [
        ("Master of Education, Educational Leadership",
         "Lamar University, Beaumont, TX (2010)"),
        ("Master of Divinity with Biblical Languages",
         "Southwestern Baptist Theological Seminary, Fort Worth, TX (2002)"),
        ("BS, Occupational Education — Russian, Religion & Crypto-Linguistics",
         "Wayland Baptist University, San Antonio, TX (1999)"),
    ]
    for deg, school in edu_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        d_run = p.add_run(deg + "  —  ")
        d_run.bold = True
        d_run.font.size = Pt(10)
        d_run.font.name = "Arial"
        _set_color(d_run, BLACK)
        s_run = p.add_run(school)
        s_run.font.size = Pt(10)
        s_run.font.name = "Arial"
        _set_color(s_run, BLACK)

    _hr(doc)

    # ── Military ─────────────────────────────────────────────────────────────
    _section_heading(doc, "Military Experience")
    mil_para = doc.add_paragraph()
    mil_para.paragraph_format.space_before = Pt(2)
    mil_run = mil_para.add_run(
        "US Air Force — Airborne Russian Cryptologic Linguist, Worldwide (1983–1999)"
    )
    mil_run.font.size = Pt(10)
    mil_run.font.name = "Arial"
    _set_color(mil_run, BLACK)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
